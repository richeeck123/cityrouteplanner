import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set the internal margins (padding) of a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Inter'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700

    # 1. Document Title / Cover Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("CITY ROUTE PLANNER")
    run_title.font.name = 'Inter'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a) # Deep blue

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("Technical Documentation, Folder Structure, Algorithms & Line-by-Line Code Report")
    run_sub.font.name = 'Inter'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69) # Slate 600

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(48)
    run_meta = meta_p.add_run("Created for: City Route Planner Project\nDate: June 2026\nStatus: Final Production")
    run_meta.font.name = 'Inter'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8b) # Slate 500

    doc.add_page_break()

    # Section 1
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Executive Summary & Folder Structure")
    run_h1.font.name = 'Inter'
    run_h1.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    run_h1.font.bold = True
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)

    p1 = doc.add_paragraph()
    p1.add_run(
        "The City Route Planner is a high-performance C++ backend application coupled with a clean HTML/JS "
        "frontend. Its core objective is to load OpenStreetMap (OSM) XML data, parse it to construct a weighted "
        "directed/undirected graph representation of a city's road network, run shortest-path routing algorithms "
        "(specifically comparing Dijkstra's Algorithm and A* Search), benchmark their execution speed and efficiency "
        "(in terms of nodes explored), export the path to JSON format, and visualize the path in a modern, interactive web dashboard."
    )
    p1.paragraph_format.space_after = Pt(12)

    h2 = doc.add_heading(level=2)
    run_h2 = h2.add_run("1.1. Folder Structure and Contents")
    run_h2.font.name = 'Inter'
    run_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph("Below is an overview of the directory structure layout and the specific purpose of each file and folder:")
    p2.paragraph_format.space_after = Pt(12)

    # Table of folders
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Folder / File Path'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Purpose & Work'

    # Set background for header
    for cell in hdr_cells:
        set_cell_background(cell, "1e3a8a")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        # Format text to white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    folders_data = [
        ("src/", "Directory", "Contains the core C++ codebase including graph structures, parser utility, search algorithms, benchmarks, and the main executable entry point."),
        ("src/Graph.h", "File", "Defines the custom Node struct, the Haversine distance calculator, and the Graph class with nodes map and weighted adjacency list storage."),
        ("src/Parser.h", "File", "Provides parsing code for .osm XML formats, dynamically extracting tag coordinates and building graph edges."),
        ("src/Algorithms.h", "File", "Implements the routing algorithms: Dijkstra's Algorithm, A* Search, Breadth-First Search (BFS), and the linear Nearest Node search."),
        ("src/Benchmark.h", "File", "Implements runtime profiling using C++ high_resolution_clock, logging time in microseconds and nodes visited for Dijkstra and A*."),
        ("src/main.cpp", "File", "The driver class representing the entry point of the app. Interacts with the user via console inputs, manages algorithm execution, and serializes path output."),
        ("data/", "Directory", "Stores the input raw mapping datasets. Holds city.osm (representing real spatial nodes and highway paths)."),
        ("images/", "Directory", "Houses demo screenshots and graphics representing path visualizations for documentation."),
        ("output/", "Directory", "Workspace directory containing output files like path_output.json containing path coordinates for the frontend."),
        ("map.html", "File", "The upgraded frontend interface. Loads Leaflet.js, renders route coordinates, toggle themes, and provides click-to-copy coordinate picks."),
        ("README.md", "File", "Comprehensive project setup instructions, compilation guidelines, viewing information, and benchmarking summaries.")
    ]

    for path, ftype, desc in folders_data:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = ftype
        row_cells[2].text = desc
        
        # Style row cells
        for cell in row_cells:
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
        # Add bolding to path column
        for run in row_cells[0].paragraphs[0].runs:
            run.font.bold = True

    doc.add_page_break()

    # Section 2
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("2. Core Algorithms Explained in Detail")
    run_h1.font.name = 'Inter'
    run_h1.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)

    # 2.1 Haversine
    h2 = doc.add_heading(level=2)
    run_h2 = h2.add_run("2.1. Haversine Formula")
    run_h2.font.name = 'Inter'
    run_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run("The Haversine formula determines the great-circle distance between two points on a sphere given their longitudes and latitudes. In City Route Planner, it is the mathematical foundation for calculating physical distance weights between nodes, and acts as the heuristic for A* Search. The equation is:")
    p.paragraph_format.space_after = Pt(6)
    
    math_p = doc.add_paragraph()
    math_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    math_run = math_p.add_run(
        "d = 2 * R * arcsin( sqrt( sin²(Δlat/2) + cos(lat1)*cos(lat2)*sin²(Δlon/2) ) )\n"
        "Where R = 6371.0 km (Earth mean radius), lat/lon are in radians."
    )
    math_run.font.bold = True
    math_run.font.name = 'Consolas'
    math_run.font.size = Pt(10)
    math_p.paragraph_format.space_after = Pt(12)

    # 2.2 Dijkstra
    h2 = doc.add_heading(level=2)
    run_h2 = h2.add_run("2.2. Dijkstra's Algorithm")
    run_h2.font.name = 'Inter'
    run_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run(
        "Dijkstra's Algorithm is a classic graph search algorithm that finds the shortest paths between nodes in a graph. "
        "It maintains a set of visited nodes and a map of tentative distances (initialized to infinity, and 0 for the source node). "
        "At each step, it selects the unvisited node with the lowest tentative distance, relaxes its outgoing edges, and updates "
        "their tentative distances. A priority queue (min-heap) is used to efficiently retrieve the minimum-distance node. "
        "Its Time Complexity is O((V + E) log V), where V is the number of vertices and E is the number of edges. "
        "Its Space Complexity is O(V) to store distances and parent tracking nodes."
    )
    p.paragraph_format.space_after = Pt(12)

    # 2.3 A* Search
    h2 = doc.add_heading(level=2)
    run_h2 = h2.add_run("2.3. A* Search Algorithm")
    run_h2.font.name = 'Inter'
    run_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run(
        "A* Search is an extension of Dijkstra's Algorithm that uses a heuristic function to guide its search. "
        "Instead of ordering nodes strictly by distance from the source g(n), it orders them by f(n) = g(n) + h(n), where "
        "h(n) is the heuristic representing the estimated cost from node n to the destination. "
        "In this project, h(n) is computed using the Haversine distance, which is an admissible heuristic (it never overestimates "
        "the actual straight-line distance on Earth's surface) and consistent, ensuring A* finds the optimal path. "
        "This heuristic drastically reduces the search space, focusing the exploration directly toward the target destination. "
        "Time Complexity is O((V + E) log V) in the worst case (equal to Dijkstra's), but practically runs much faster (often 5-10x "
        "fewer nodes explored) depending on map topology."
    )
    p.paragraph_format.space_after = Pt(12)

    # 2.4 BFS
    h2 = doc.add_heading(level=2)
    run_h2 = h2.add_run("2.4. Breadth-First Search (BFS)")
    run_h2.font.name = 'Inter'
    run_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run(
        "BFS is a fundamental unweighted graph traversal algorithm. It explores all neighbor nodes at the current depth level "
        "before moving to nodes at the next level, utilizing a FIFO queue. In the context of weighted road networks, BFS is "
        "sub-optimal because it assumes all road weights are equal to 1, completely ignoring actual geographical distances. "
        "It serves as a baseline comparison. Time Complexity is O(V + E), and Space Complexity is O(V)."
    )
    p.paragraph_format.space_after = Pt(12)

    # 2.5 Nearest Node Search
    h2 = doc.add_heading(level=2)
    run_h2 = h2.add_run("2.5. Nearest Node Search")
    run_h2.font.name = 'Inter'
    run_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.add_run(
        "To resolve user-input GPS coordinates (latitude and longitude) to actual vertices in the parsed graph, "
        "the application performs a nearest-node search. It iterates through all available nodes, calculates the Haversine "
        "distance between the input coordinates and the node coordinates, and returns the node ID that minimizes this distance. "
        "Currently, this linear scan runs in O(V) time. It can be optimized to O(log V) in the future using spatial index trees "
        "(such as KD-Trees or R-Trees)."
    )
    p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Section 3
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("3. Line-by-Line Code Walkthrough & Report")
    run_h1.font.name = 'Inter'
    run_h1.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)

    # Walkthrough list
    walkthrough_files = [
        ("src/Graph.h", "Defines nodes and structural connections.", [
            ("Lines 1-6: Directives & Namespaces", 
             "#pragma once handles header guard compilation safety. Unordered maps, vectors, and cmath headers are imported. namespace std is used for ease."),
            ("Lines 8-10: Pi constant definition",
             "Defines M_PI mathematically if it is not already declared in the local environment, used for degree-to-radian coordinates conversion."),
            ("Lines 12-15: Node definition struct",
             "Defines a GPS point node with a long long id, double lat, and double lon."),
            ("Lines 17-25: haversine() calculator",
             "Calculates spatial distance using Earth's sphere radius (6371.0 km) and basic trigonometry. Degrees are multiplied by M_PI/180.0 to convert to radians. The arcsin formulation is calculated using standard sqrt and atan2."),
            ("Lines 27-31: Graph class and attributes",
             "Graph manages a map of node IDs to Node structs (`nodes`) and a map of node IDs to list of weighted adjacency pairs `adj` (representing neighbor node ID and double weight)."),
            ("Lines 32-34: addNode() method",
             "Creates a node object in the map, lookup is key O(1)."),
            ("Lines 36-46: addEdge() method",
             "Accepts two node IDs, verifies both exist in the `nodes` lookup, computes the Haversine distance, and adds the connection symmetrically to build an undirected road graph.")
        ]),
        ("src/Parser.h", "Responsible for parsing OpenStreetMap XML data files.", [
            ("Lines 1-8: Imports and includes",
             "Includes Graph.h, fstream, string, and vector. Standard std namespace is activated."),
            ("Lines 9-14: extractLongLong() parsing tool",
             "Locates substring matching the attribute name (e.g. id=\" or ref=\"), identifies enclosing quotes, extracts the inner characters, and converts them to a long long value via stoll."),
            ("Lines 16-21: extractDouble() parsing tool",
             "Similar to extractLongLong but converts values to floating-point doubles using stod for lat/lon fields."),
            ("Lines 23-28: parseOSM() header and initialization",
             "Loads filename into std::ifstream and prepares wayNodes vector to store sequential nodes describing a road segment. isHighway tracks road flags."),
            ("Lines 29-35: XML node parsing",
             "Iteratively reads lines. If <node tag is matched, parses ID, lat, and lon attributes using utility extractors and updates the graph object."),
            ("Lines 37-47: XML way elements",
             "If <way tag starts, clears temporary node lists. If <nd ref= tag is found, stores reference IDs. If k=\"highway\" is found, isHighway flag is set to true. Roads are categorized as highways in OSM format."),
            ("Lines 48-54: Way finalization",
             "On closing tag </way>, if the segment is validated as a highway, iterates through wayNodes array adding bidirectional graph edges between sequential vertices (e.g., node i to node i+1).")
        ]),
        ("src/Algorithms.h", "Contains all path search algorithms.", [
            ("Lines 11-17: dijkstra() function signature",
             "Accepts Graph g reference, source ID, destination ID, and an optional nodesVisited count pointer. Sets up local distance tracker, previous parent map tracker, and standard priority queue (min-priority structure)."),
            ("Lines 19-23: Dijkstra initialization",
             "Sets all distances in graph nodes to infinity (represented by 1e18), initializes source distance to 0, pushes the source node into priority queue, and resets visited nodes count to 0."),
            ("Lines 25-32: Dijkstra main loop",
             "Pops node with smallest distance curDist. Increments nodes visited. If destination is dequeued, breaks early. If curDist is older/larger than tracked node distance, ignores it."),
            ("Lines 33-42: Dijkstra edge relaxation",
             "Iterates neighbors of current node. If path distance through current node is smaller than neighbor's tracked distance, updates distance, logs current node as neighbor's parent (`prev` map), and queues neighbor node."),
            ("Lines 43-44: Dijkstra finalization & error handling",
             "Stores final nodes visited count. If destination distance remains near infinity, returns failure flag and empty route array."),
            ("Lines 45-55: Path reconstruction",
             "Backtracks from destination to source using the `prev` parent map, reverses the resulting vector, and returns the sorted route path along with computed distance."),
            ("Lines 57-64: heuristic() function",
             "Computes A* heuristic. Uses haversine() to calculate direct distance between current node u and destination. Admissible and consistent."),
            ("Lines 66-78: astar() initialization",
             "Uses gScore map initialized to infinity. Sets gScore[src] to 0. Priority queue holds fScore (gScore + heuristic) and node ID."),
            ("Lines 80-96: A* main loop and relaxation",
             "Pops lowest fScore node. Breaks if destination. Relaxes neighbors using edge weight w. If new gScore is smaller than recorded gScore, updates it, logs previous parent node, and queues neighbor with updated fScore: gScore + heuristic."),
            ("Lines 98-110: A* output generation",
             "Similar parent backtracking logic to reconstruct optimal coordinate path, reversing output array and returning total distance."),
            ("Lines 112-134: bfs() implementation",
             "Implements a standard Breadth-First Search queue traversal. Tracks visited IDs in a set. Returns first route found by topological depth, ignoring weight distance values."),
            ("Lines 136-146: BFS path reconstruction",
             "Assembles and reverses node path based on parent records logged during FIFO queue processing."),
            ("Lines 148-167: nearestNode() scanner",
             "Loops through graph nodes map. If node has active neighbor connections, computes haversine distance. Tracks minimum distance and returns matching node ID.")
        ]),
        ("src/Benchmark.h", "Contains the performance benchmarking class.", [
            ("Lines 11-24: BenchResult class",
             "A class carrying result properties: distance in km, number of nodes visited, execution duration in microseconds, and vector of route node IDs."),
            ("Lines 25-34: runDijkstra() method",
             "Starts high resolution clock, executes dijkstra(), records time difference, and moves the final path vector to build the BenchResult object."),
            ("Lines 35-42: runAStar() method",
             "Profiles astar() search execution in microseconds, tracks node counts visited, and outputs metrics in BenchResult format."),
            ("Lines 43-50: printComparison() method",
             "Outputs algorithm metrics side-by-side: distance, visited nodes, time. Calculates search space reduction efficiency as a speedup factor (Dijkstra visited / A* visited).")
        ]),
        ("src/main.cpp", "The main executable CLI coordinator.", [
            ("Lines 11-30: exportPath() output routine",
             "Creates a JSON file 'output/path_output.json'. Writes algorithm name, distance in km, and coordinate list arrays representing latitude/longitude values."),
            ("Lines 32-45: main() load sequence",
             "Prints 'City Route Planner' title. Calls parseOSM. Calculates total bidirectional edges (edges sum/2) and reports node and edge counts."),
            ("Lines 46-55: User console inputs",
             "Prompts user for double decimal coordinate inputs: source latitude/longitude and destination latitude/longitude via standard cin streams."),
            ("Lines 56-64: Resolution and validation",
             "Calls nearestNode() lookup for source and destination coordinates. Validates input bounds (-1 indicates no matching active nodes)."),
            ("Lines 66-68: Run & Print comparison",
             "Executes dijkstra() and astar() benchmarks, then prints the detailed comparison statistics to console stdout."),
            ("Lines 70-81: Path export",
             "Checks if computed path is non-empty. Prefers exporting the A* path (or fallback Dijkstra path) to JSON output, and terminates normally.")
        ]),
        ("map.html", "The interactive browser frontend dashboard.", [
            ("Lines 1-22: CSS styling and variables",
             "Defines Inter font, color variables for light and dark themes, background colors, card dimensions, glassmorphism filters, badges, buttons, and transition keyframes."),
            ("Lines 24-44: HTML Layout markup",
             "Builds the side floating card control panel. Declares sections for branding/logo, active route metrics, and click-coordinate picker controls."),
            ("Lines 46-100: Map Setup and theme handling",
             "Initializes Leaflet map. Sets up Light/Dark CartoDB base maps. Configures theme toggle event listener to dynamically switch styles and map tile layers."),
            ("Lines 101-147: Path rendering logic",
             "Fetches path_output.json. Adds a shadow layer to represent the route, builds the primary route polyline (vibrant blue), creates circle markers for Start and End, and fits map view bounds to the path."),
            ("Lines 148-185: Interactive picker handler",
             "Adds map.on('click') listener to record coordinates, renders selections on screen, creates visual markers, copies text representation to clipboard on button click, and displays custom toast alerts.")
        ])
    ]

    for title, desc, lines in walkthrough_files:
        h2 = doc.add_heading(level=2)
        run_h2 = h2.add_run(f"3.x. Walkthrough for {title}")
        run_h2.font.name = 'Inter'
        run_h2.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(4)

        p = doc.add_paragraph()
        p.add_run(f"File purpose: {desc}")
        p.paragraph_format.space_after = Pt(8)

        # Walkthrough details in table
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Light Shading Accent 1'
        hdr = t.rows[0].cells
        hdr[0].text = 'Line Range / Block'
        hdr[1].text = 'Detailed Logical Work / Action'
        
        for c in hdr:
            set_cell_background(c, "2563eb")
            set_cell_margins(c, top=100, bottom=100, left=120, right=120)
            for p_cell in c.paragraphs:
                for r in p_cell.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                    r.font.size = Pt(9.5)

        for line_range, explanation in lines:
            row = t.add_row().cells
            row[0].text = line_range
            row[1].text = explanation
            
            for c in row:
                set_cell_margins(c, top=80, bottom=80, left=100, right=100)
                for p_cell in c.paragraphs:
                    for r in p_cell.runs:
                        r.font.size = Pt(9.0)
            
            # Make line range bold
            for r in row[0].paragraphs[0].runs:
                r.font.bold = True

        doc.add_page_break()

    # Save document
    doc.save("City_Route_Planner_Report.docx")
    print("Report generated successfully as 'City_Route_Planner_Report.docx'")

if __name__ == "__main__":
    create_report()
